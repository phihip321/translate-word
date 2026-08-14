import socket
from docx import Document
import tkinter as tk
from tkinter import filedialog
from dotenv import load_dotenv
from google import genai

# IPv4
old_getaddrinfo = socket.getaddrinfo

def getaddrinfo_ipv4(host, port, family=0, type=0, proto=0, flags=0):
    return old_getaddrinfo(
        host,
        port,
        socket.AF_INET,
        type,
        proto,
        flags
    )

socket.getaddrinfo = getaddrinfo_ipv4

# Gemini
load_dotenv()
client = genai.Client()


# Hàm dịch một nhóm đoạn
def translate_batch(text):

    prompt = f"""
Dịch các đoạn văn tiếng Anh sau sang tiếng Việt.

Yêu cầu:
- Giữ đúng thứ tự các đoạn.
- Mỗi đoạn bắt đầu bằng [ĐOẠN].
- Bản dịch cũng phải bắt đầu bằng [ĐOẠN].
- Không giải thích.
- Giữ nguyên thuật ngữ y khoa.

Nội dung:

{text}
"""

    interaction = client.interactions.create(
        model="gemini-3.5-flash",
        input=prompt,
        generation_config={
            "thinking_level": "minimal"
        }
    )

    return interaction.output_text


# Chọn file
root = tk.Tk()
root.withdraw()

file_path = filedialog.askopenfilename(
    title="Chọn file Word",
    filetypes=[("Word Document", "*.docx")]
)

root.destroy()

if not file_path:
    print("Bạn chưa chọn file.")
    exit()


# Đọc Word
doc = Document(file_path)

paragraphs = []

for paragraph in doc.paragraphs:
    text = paragraph.text.strip()

    if text:
        paragraphs.append(text)


# Chia thành nhóm
batch_size = 10

translated_all = []

print()
print(f"Tổng số đoạn: {len(paragraphs)}")
print()


for i in range(0, len(paragraphs), batch_size):

    batch = paragraphs[i:i + batch_size]

    numbered_text = ""

    for text in batch:
        numbered_text += f"[ĐOẠN]\n{text}\n\n"

    print(
        f"Đang dịch nhóm "
        f"{i + 1} - {i + len(batch)}..."
    )

    result = translate_batch(numbered_text)

    translated_all.append(result)


# Tạo Word mới
new_doc = Document()

for result in translated_all:

    parts = result.split("[ĐOẠN]")

    for part in parts:

        text = part.strip()

        if text:
            new_doc.add_paragraph(text)


# Lưu
output_file = file_path.rsplit(".", 1)[0] + "_VI.docx"

new_doc.save(output_file)

print()
print("===== HOÀN THÀNH =====")
print(output_file)