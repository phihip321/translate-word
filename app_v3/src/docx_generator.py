"""
Tạo file Word từ bản dịch
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List
from pathlib import Path

class DOCXGenerator:
    def __init__(self, config: dict):
        self.config = config
    
    def create_translation_document(self,
                                   original_paragraphs: List[str],
                                   translated_paragraphs: List[str],
                                   chapter_name: str,
                                   output_path: str) -> str:
        """
        Tạo file Word từ bản dịch
        """
        doc = Document()
        
        # Thiết lập margin
        for section in doc.sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
        
        # Thêm tiêu đề
        title = doc.add_heading(chapter_name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ngăn cách
        doc.add_paragraph('_' * 50)
        
        # Duyệt qua từng paragraph
        for i, (orig, trans) in enumerate(zip(original_paragraphs, translated_paragraphs), 1):
            # Đánh số paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            
            # Thêm số thứ tự
            run = p.add_run(f"{i}. ")
            run.bold = True
            run.font.size = Pt(self.config.get('font_size', 11))
            
            if self.config.get('show_original', True):
                # Bản gốc (tiếng Anh)
                run = p.add_run(orig)
                run.font.name = self.config.get('font_name', 'Times New Roman')
                run.font.size = Pt(self.config.get('font_size', 11))
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)  # Màu xám
                
                p.add_run('\n')
            
            if self.config.get('show_translation', True):
                # Bản dịch (tiếng Việt)
                run = p.add_run(trans)
                run.font.name = self.config.get('font_name', 'Times New Roman')
                run.font.size = Pt(self.config.get('font_size', 11))
                run.font.bold = True
            
            # Khoảng cách dòng
            p.paragraph_format.line_spacing = self.config.get('line_spacing', 1.5)
            p.paragraph_format.space_after = Pt(6)
        
        # Lưu file
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        return str(output_path)
    
    def create_parallel_document(self,
                                original_paragraphs: List[str],
                                translated_paragraphs: List[str],
                                chapter_name: str,
                                output_path: str) -> str:
        """
        Tạo file Word 2 cột (Anh - Việt song song)
        """
        doc = Document()
        
        # Margin rộng hơn cho 2 cột
        for section in doc.sections:
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
        
        # Tiêu đề
        title = doc.add_heading(chapter_name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tạo bảng 2 cột
        table = doc.add_table(rows=len(original_paragraphs), cols=2)
        table.style = 'Table Grid'
        
        # Điều chỉnh độ rộng cột (50-50)
        for row in table.rows:
            row.cells[0].width = Inches(4.5)
            row.cells[1].width = Inches(4.5)
        
        # Đổ dữ liệu
        for i, (orig, trans) in enumerate(zip(original_paragraphs, translated_paragraphs)):
            # Cột tiếng Anh
            cell = table.cell(i, 0)
            cell.text = f"{i+1}. {orig}"
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.italic = True
            
            # Cột tiếng Việt
            cell = table.cell(i, 1)
            cell.text = f"{i+1}. {trans}"
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.bold = True
            
        # Lưu
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        return str(output_path)