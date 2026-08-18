"""
Trích xuất PDF thành Word với:
1. GIỮ NGUYÊN FORMATTING TEXT (font, size, bold, italic)
2. Hình ảnh chèn đúng vị trí
3. Bảng biểu
4. Loại bỏ header/footer
"""

import os
import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
from PIL import Image
import io

# ============================================================================
# CẤU HÌNH
# ============================================================================
OUTPUT_BASE_DIR = r"D:\python\translate-word\output"
HEADER_RATIO = 0.08  # Cắt 8% top
FOOTER_RATIO = 0.08  # Cắt 8% bottom

# ============================================================================
# HÀM TIỆN ÍCH
# ============================================================================
def is_header_or_footer(bbox, page_height, header_ratio=HEADER_RATIO, footer_ratio=FOOTER_RATIO):
    """Kiểm tra block có phải header/footer không"""
    y0, y1 = bbox[1], bbox[3]
    top_boundary = page_height * header_ratio
    bottom_boundary = page_height * (1 - footer_ratio)
    return y1 <= top_boundary or y0 >= bottom_boundary

def set_run_formatting(run, span):
    """
    Copy formatting từ PDF span sang Word run
    """
    # Font name
    font_name = span.get("font", "")
    if font_name:
        # Map font từ PDF sang font phổ biến
        font_mapping = {
            "timesnewroman": "Times New Roman",
            "times": "Times New Roman",
            "arial": "Arial",
            "helvetica": "Arial",
            "calibri": "Calibri",
            "cambria": "Cambria",
        }
        
        font_lower = font_name.lower()
        for key, value in font_mapping.items():
            if key in font_lower:
                run.font.name = value
                break
    
    # Font size
    font_size = span.get("size", 12)
    if font_size:
        run.font.size = Pt(font_size)
    
    # Bold
    flags = span.get("flags", 0)
    if flags & 2**4 or "bold" in font_name.lower() or "black" in font_name.lower():
        run.bold = True
    
    # Italic
    if flags & 2**1 or "italic" in font_name.lower() or "oblique" in font_name.lower():
        run.italic = True
    
    # Underline
    if flags & 2**2:
        run.underline = True
    
    # Color
    color = span.get("color", 0)
    if color:
        # PDF color là int (RGB)
        r = (color >> 16) & 0xFF
        g = (color >> 8) & 0xFF
        b = color & 0xFF
        
        # Chỉ set màu nếu không phải đen (0,0,0)
        if not (r == 0 and g == 0 and b == 0):
            run.font.color.rgb = RGBColor(r, g, b)
    
    # Superscript/Subscript
    if flags & 2**0:  # Superscript
        run.font.superscript = True
    if flags & 2**5:  # Subscript
        run.font.subscript = True

def get_span_text_with_formatting(span):
    """
    Lấy text và thông tin formatting từ span
    """
    return {
        'text': span.get("text", ""),
        'font': span.get("font", ""),
        'size': span.get("size", 12),
        'flags': span.get("flags", 0),
        'color': span.get("color", 0),
    }

def process_pdf_to_formatted_docx():
    """Trích xuất PDF thành Word giữ nguyên formatting"""
    
    # 1. Mở cửa sổ chọn file PDF
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)

    pdf_path = filedialog.askopenfilename(
        title="Chọn file PDF cần trích xuất (Giữ nguyên formatting)",
        filetypes=[("PDF Files", "*.pdf")]
    )

    if not pdf_path:
        print("Bạn chưa chọn file nào.")
        return

    pdf_doc = fitz.open(pdf_path)
    docx_doc = Document()
    
    # Thiết lập style mặc định
    style = docx_doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # 2. Tạo cấu trúc thư mục
    base_name = os.path.basename(pdf_path)
    file_name, _ = os.path.splitext(base_name)
    
    book_output_dir = os.path.join(OUTPUT_BASE_DIR, file_name)
    images_dir = os.path.join(book_output_dir, "images")
    
    os.makedirs(images_dir, exist_ok=True)
    
    output_path = os.path.join(book_output_dir, f"{file_name}_formatted.docx")

    print(f"📄 Đang xử lý file: {base_name}")
    print(f"🎨 Giữ nguyên formatting text...")
    
    # 3. Xử lý từng trang
    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        height = page.rect.height
        width = page.rect.width
        mid_x = width / 2
        
        # === A. PHÁT HIỆN BẢNG ===
        tables = page.find_tables()
        table_bboxes = [t.bbox for t in tables]
        
        # === B. LẤY TEXT VÀ IMAGE BLOCKS ===
        page_dict = page.get_text("dict")
        blocks = page_dict.get("blocks", [])
        
        valid_blocks = []
        for b in blocks:
            bbox = b["bbox"]
            
            # Loại bỏ header/footer
            if is_header_or_footer(bbox, height):
                continue
            
            # Loại bỏ blocks nằm trong bảng
            in_table = False
            for t_bbox in table_bboxes:
                if (bbox[0] >= t_bbox[0] and bbox[1] >= t_bbox[1] and 
                    bbox[2] <= t_bbox[2] and bbox[3] <= t_bbox[3]):
                    in_table = True
                    break
            
            if not in_table:
                valid_blocks.append(b)
        
        # === C. SẮP XẾP THEO CỘT ===
        left_blocks = [b for b in valid_blocks if b["bbox"][2] <= mid_x + 15]
        right_blocks = [b for b in valid_blocks if b["bbox"][0] >= mid_x - 15]
        
        is_two_col = len(left_blocks) >= 2 and len(right_blocks) >= 2
        
        if is_two_col:
            sorted_blocks = sorted(left_blocks, key=lambda x: x["bbox"][1]) + \
                            sorted(right_blocks, key=lambda x: x["bbox"][1])
        else:
            sorted_blocks = sorted(valid_blocks, key=lambda x: x["bbox"][1])
        
        # === D. XUẤT TEXT VỚI FORMATTING ===
        for b_idx, b in enumerate(sorted_blocks):
            if b.get("type") == 0:  # Text block
                # Lấy tất cả lines và spans
                for line in b.get("lines", []):
                    # Kiểm tra line có text không
                    line_text = ""
                    for span in line.get("spans", []):
                        line_text += span.get("text", "")
                    
                    if not line_text.strip():
                        continue
                    
                    # Tạo paragraph cho line
                    p = docx_doc.add_paragraph()
                    
                    # Thêm từng span với formatting
                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        if not text:
                            continue
                        
                        # Tạo run
                        run = p.add_run(text)
                        
                        # Copy formatting
                        set_run_formatting(run, span)
                    
                    # === ĐỊNH DẠNG PARAGRAPH ===
                    
                    # Line spacing
                    p.paragraph_format.line_spacing = 1.15
                    
                    # Space after
                    p.paragraph_format.space_after = Pt(4)
                    
                    # Alignment (dựa trên vị trí text)
                    bbox = b["bbox"]
                    x_center = (bbox[0] + bbox[2]) / 2
                    
                    if x_center < mid_x * 0.4:
                        # Text ở bên trái
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif x_center > mid_x * 1.6:
                        # Text ở bên phải
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif abs(x_center - mid_x) < mid_x * 0.15:
                        # Text ở giữa
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        # Mặc định
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    # First line indent (nếu text thụt vào)
                    first_span_x = bbox[0]
                    if first_span_x > width * 0.15:  # Thụt vào >15% trang
                        p.paragraph_format.first_line_indent = Cm(1.27)  # ~0.5 inch
            
            elif b.get("type") == 1:  # Image block
                try:
                    img_bytes = b.get("image")
                    img_ext = b.get("ext", "png")
                    
                    if not img_bytes or len(img_bytes) < 100:
                        continue
                    
                    img_code = f"IMAGE_p{page_num + 1}_b{b_idx}"
                    temp_img_path = os.path.join(images_dir, f"{img_code}.{img_ext}")
                    
                    # Xử lý ảnh
                    try:
                        image_stream = io.BytesIO(img_bytes)
                        pil_img = Image.open(image_stream)
                        
                        if pil_img.mode in ("CMYK", "P"):
                            pil_img = pil_img.convert("RGB")
                        
                        pil_img.save(temp_img_path)
                    except Exception:
                        with open(temp_img_path, "wb") as f:
                            f.write(img_bytes)
                    
                    # Chèn ảnh vào Word
                    p = docx_doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Chèn ảnh với kích thước phù hợp
                    try:
                        # Lấy kích thước ảnh
                        with Image.open(temp_img_path) as img:
                            img_w, img_h = img.size
                            
                            # Scale ảnh để vừa trang Word (max 5 inch)
                            max_width = 5.0
                            if img_w > 0:
                                ratio = max_width / img_w
                                display_width = Inches(max_width)
                                display_height = Inches(img_h * ratio)
                                
                                # Nếu ảnh quá cao, giới hạn chiều cao
                                if display_height.inches > 7.0:
                                    display_height = Inches(7.0)
                                    display_width = Inches(7.0 * img_w / img_h)
                                
                                p.add_run().add_picture(temp_img_path, width=display_width)
                    except:
                        # Fallback: chèn với width mặc định
                        p.add_run().add_picture(temp_img_path, width=Inches(5.0))
                    
                    # Thêm khoảng trắng sau ảnh
                    p.paragraph_format.space_after = Pt(6)
                    
                except Exception as e:
                    print(f"   ⚠️ Bỏ qua ảnh lỗi ở trang {page_num + 1} (Block {b_idx}): {e}")
        
        # === E. XUẤT BẢNG VỚI FORMATTING ===
        for table_idx, table in enumerate(tables):
            if is_header_or_footer(table.bbox, height):
                continue
            
            data = table.extract()
            if not data:
                continue
            
            # Thêm khoảng trắng trước bảng
            docx_doc.add_paragraph()
            
            # Tạo bảng
            doc_table = docx_doc.add_table(rows=len(data), cols=len(data[0]))
            doc_table.style = "Table Grid"
            
            # Format từng cell
            for r_idx, row in enumerate(data):
                for c_idx, cell_value in enumerate(row):
                    cell = doc_table.cell(r_idx, c_idx)
                    cell.text = cell_value or ""
                    
                    # Format paragraph trong cell
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(2)
                        paragraph.paragraph_format.space_before = Pt(2)
                        
                        # Header row - bold
                        if r_idx == 0:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(11)
                        else:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
            
            # Thêm khoảng trắng sau bảng
            docx_doc.add_paragraph()
        
        # Hiển thị tiến trình
        if (page_num + 1) % 5 == 0 or page_num == len(pdf_doc) - 1:
            print(f"   ✅ Đã xử lý {page_num + 1}/{len(pdf_doc)} trang")
    
    # 4. Lưu file
    docx_doc.save(output_path)
    
    print(f"\n{'='*80}")
    print(f"✅ HOÀN THÀNH!")
    print(f"{'='*80}")
    print(f"📁 File Word: {output_path}")
    print(f"📁 Thư mục ảnh: {images_dir}")
    print(f"📊 Tổng số trang: {len(pdf_doc)}")
    print(f"🎨 Định dạng: Giữ nguyên font, size, bold, italic")
    
    # Đóng PDF
    pdf_doc.close()
    
    # Mở thư mục
    try:
        os.startfile(book_output_dir)
    except Exception as e:
        print(f"⚠️ Không thể mở thư mục: {e}")

# ============================================================================
# MAIN
# ============================================================================
if __name__ == "__main__":
    process_pdf_to_formatted_docx()