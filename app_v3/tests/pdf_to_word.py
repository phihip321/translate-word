import os
import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF
from PIL import Image, ImageOps, UnidentifiedImageError
import io

def is_header_or_footer(bbox, page_height, header_ratio=0.08, footer_ratio=0.08):
    y0, y1 = bbox[1], bbox[3]
    top_boundary = page_height * header_ratio
    bottom_boundary = page_height * (1 - footer_ratio)
    return y1 <= top_boundary or y0 >= bottom_boundary

def process_pdf_to_formatted_docx():
    # 1. Mở cửa sổ chọn file PDF
    root = tk.Tk()
    root.withdraw()

    pdf_path = filedialog.askopenfilename(
        title="Chọn file PDF cần trích xuất (Formatted Text)",
        filetypes=[("PDF Files", "*.pdf")]
    )

    if not pdf_path:
        print("Bạn chưa chọn file nào.")
        return

    pdf_doc = fitz.open(pdf_path)
    docx_doc = Document()

    # 2. Cấu trúc thư mục phân tầng cho từng cuốn sách
    base_name = os.path.basename(pdf_path)
    file_name, _ = os.path.splitext(base_name)
    
    # Tạo thư mục riêng: output/ten_file_sach/
    book_output_dir = os.path.join(r"D:\python\translate-word\output", file_name)
    # Tạo thư mục chứa ảnh riêng: output/ten_file_sach/images/
    images_dir = os.path.join(book_output_dir, "images")
    
    os.makedirs(images_dir, exist_ok=True)
    
    output_path = os.path.join(book_output_dir, f"{file_name}_formatted.docx")

    print(f"Đang xử lý file: {base_name} ...")

    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        height = page.rect.height
        width = page.rect.width
        mid_x = width / 2

        # A. Bảng (Tables)
        tables = page.find_tables()
        table_bboxes = [t.bbox for t in tables]

        # B. Lấy Text và Image
        page_dict = page.get_text("dict")
        blocks = page_dict.get("blocks", [])

        valid_blocks = []
        for b in blocks:
            bbox = b["bbox"]

            if is_header_or_footer(bbox, height):
                continue

            in_table = False
            for t_bbox in table_bboxes:
                if (bbox[0] >= t_bbox[0] and bbox[1] >= t_bbox[1] and 
                    bbox[2] <= t_bbox[2] and bbox[3] <= t_bbox[3]):
                    in_table = True
                    break

            if not in_table:
                valid_blocks.append(b)

        # C. Sắp xếp Cột (Trái -> Phải)
        left_blocks = [b for b in valid_blocks if b["bbox"][2] <= mid_x + 15]
        right_blocks = [b for b in valid_blocks if b["bbox"][0] >= mid_x - 15]

        is_two_col = len(left_blocks) >= 2 and len(right_blocks) >= 2

        if is_two_col:
            sorted_blocks = sorted(left_blocks, key=lambda x: x["bbox"][1]) + \
                            sorted(right_blocks, key=lambda x: x["bbox"][1])
        else:
            sorted_blocks = sorted(valid_blocks, key=lambda x: x["bbox"][1])

        # D. Xuất Text và Xử lý Hình ảnh
        for b_idx, b in enumerate(sorted_blocks):
            if b.get("type") == 0:  # Khối văn bản
                p = docx_doc.add_paragraph()
                for line in b.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        if not text:
                            continue

                        run = p.add_run(text)
                        flags = span.get("flags", 0)
                        font_name = span.get("font", "").lower()

                        if flags & 2**4 or "bold" in font_name or "black" in font_name:
                            run.bold = True
                        if flags & 2**1 or "italic" in font_name or "oblique" in font_name:
                            run.italic = True

            elif b.get("type") == 1:  # Khối hình ảnh
                try:
                    img_bytes = b.get("image")
                    img_ext = b.get("ext", "png")
                    
                    if not img_bytes or len(img_bytes) < 100:
                        continue

                    img_code = f"IMAGE_p{page_num + 1}_b{b_idx}"
                    temp_img_path = os.path.join(images_dir, f"{img_code}.{img_ext}")

                    # Xử lý chống lỗi âm bản/sai màu bằng Pillow
                    try:
                        image_stream = io.BytesIO(img_bytes)
                        pil_img = Image.open(image_stream)
                        
                        # Nếu ảnh dạng Mode 'CMYK' hoặc 'P' có thể chuyển về RGB để hiển thị chuẩn
                        if pil_img.mode in ("CMYK", "P"):
                            pil_img = pil_img.convert("RGB")
                            
                        # Lưu ảnh đã xử lý màu chuẩn xuống ổ cứng
                        pil_img.save(temp_img_path)
                    except Exception:
                        # Fallback: Nếu Pillow lỗi, lưu trực tiếp bytes thô
                        with open(temp_img_path, "wb") as f:
                            f.write(img_bytes)

                    # Chèn thẻ định danh và hình ảnh vào file Word
                    p = docx_doc.add_paragraph()
                    tag_run = p.add_run(f"[{img_code}]")
                    tag_run.font.color.rgb = fitz.utils.getColor("gray") if hasattr(fitz, 'utils') else None
                    
                    p.add_run().add_picture(temp_img_path, width=Inches(5.0))

                except Exception as e:
                    print(f"Bỏ qua ảnh lỗi ở trang {page_num + 1} (Block {b_idx}): {e}")

        # E. Bảng biểu (Tables)
        for table in tables:
            if is_header_or_footer(table.bbox, height):
                continue

            data = table.extract()
            if not data:
                continue

            doc_table = docx_doc.add_table(rows=len(data), cols=len(data[0]))
            doc_table.style = "Table Grid"

            for r_idx, row in enumerate(data):
                for c_idx, cell_value in enumerate(row):
                    doc_table.cell(r_idx, c_idx).text = cell_value or ""

    docx_doc.save(output_path)
    print(f"\nHoàn tất! Toàn bộ file và thư mục ảnh được lưu tại: {book_output_dir}")

    # Tự động mở thư mục riêng của sách
    try:
        os.startfile(book_output_dir)
    except Exception as e:
        print(f"Không thể mở thư mục tự động: {e}")

if __name__ == "__main__":
    process_pdf_to_formatted_docx()